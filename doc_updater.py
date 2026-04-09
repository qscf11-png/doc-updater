import pandas as pd
from docx import Document
import os
import openpyxl

def get_obsolete_list(excel_path):
    wb = openpyxl.load_workbook(excel_path, data_only=True)
    sheet = wb.worksheets[4] # 三階文件
    obsolete = []
    # 根據顏色分析，灰色是 color index 0 或 2 (在此特定檔案環境下)
    # 以及使用者截圖明確指出的範圍
    for row in sheet.iter_rows(min_row=1, max_row=200):
        cell = row[5] # F column
        if cell.fill and (str(cell.fill.start_color.index) in ['0', '2']):
            val = str(cell.value).strip()
            if val and val != 'None':
                obsolete.append(val)
                # 也加入純 ID 匹配
                if '_' in val:
                    obsolete.append(val.split('_')[0].strip())
    return list(set(obsolete))

def create_mapping(excel_path):
    xl = pd.ExcelFile(excel_path)
    sheet_name = xl.sheet_names[4]
    df = pd.read_excel(xl, sheet_name=sheet_name, header=None).fillna('')
    
    mapping = {}
    for i in range(1, len(df)):
        new_val = str(df.iloc[i, 1]).strip()
        old_val = str(df.iloc[i, 5]).strip()
        
        # 只有當兩者不同且 F 欄不是報廢(灰色)時才進行正常更換
        # (這裡我們先做正常映射，後面再套用刪除邏輯)
        if old_val and new_val and old_val != new_val:
            old_clean = old_val.replace('\n', ' ')
            mapping[old_clean] = new_val
            
            if '_' in old_clean:
                old_id = old_clean.split('_')[0].strip()
                new_id = new_val.split('_')[0].strip()
                if old_id and new_id and old_id != new_id:
                    mapping[old_id] = new_id
            
            old_name_only = old_clean.split('_', 1)[-1].strip() if '_' in old_clean else old_clean
            new_name_only = new_val.split('_', 1)[-1].strip() if '_' in new_val else new_val
            if old_name_only and new_name_only and old_name_only != new_name_only:
                mapping[old_name_only] = new_name_only

    mapping["操作指導書"] = "作業指導書"
    return mapping

def process_docx(doc_path, mapping, obsolete_list, output_path):
    doc = Document(doc_path)
    stats = {"paragraphs": 0, "tables": 0, "obsolete_removed": 0}

    def clean_and_replace(text):
        if not text: return text
        
        # 首先檢查是否是報廢文件內容 (依據 ID 或 完整名稱)
        # 我們將比對原始 ID (例如 HQ-PD.497)
        for obs in obsolete_list:
            if obs in text:
                return "" # 只要內文含有報廢 ID，就清空整體資訊
        
        # 若非報廢，則執行正常替換
        new_text = text
        for old_str, new_str in mapping.items():
            if old_str in new_text:
                new_text = new_text.replace(old_str, new_str)
        return new_text

    # 1. 段落
    for p in doc.paragraphs:
        original = p.text
        updated = clean_and_replace(original)
        if original != updated:
            p.text = updated
            if updated == "": stats["obsolete_removed"] += 1
            else: stats["paragraphs"] += 1

    # 2. 表格
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                # 注意：表格單元格可能包含多個段落
                for p in cell.paragraphs:
                    original = p.text
                    updated = clean_and_replace(original)
                    if original != updated:
                        p.text = updated
                        if updated == "": stats["obsolete_removed"] += 1
                        else: stats["tables"] += 1

    # 3. 標頭標尾
    for section in doc.sections:
        for header in [section.header, section.first_page_header, section.even_page_header]:
            if header:
                for p in header.paragraphs:
                    original = p.text
                    updated = clean_and_replace(original)
                    if original != updated:
                        p.text = updated
        for footer in [section.footer, section.first_page_footer, section.even_page_footer]:
            if footer:
                for p in footer.paragraphs:
                    original = p.text
                    updated = clean_and_replace(original)
                    if original != updated:
                        p.text = updated

    doc.save(output_path)
    return stats

if __name__ == "__main__":
    excel_file = r"temp/ISO17025(2017)實驗室管理系統文件總覽表.xlsx"
    docx_file = r"temp/template.docx" # 回歸使用原始範本，一次完成所有更換與刪除
    output_file = r"HQ-PD-P-038_v1.1_實驗室測試方法管理程序_更新版.docx"
    
    print("正在分析 Excel 顏色標示與報廢清單...")
    obs_list = get_obsolete_list(excel_file)
    print(f"成功識別 {len(obs_list)} 個報廢標示項。")
    
    print("正在提取更新映射表...")
    mapping_dict = create_mapping(excel_file)
    
    print("\n正在更新與清理 Word 文件...")
    result_stats = process_docx(docx_file, mapping_dict, obs_list, output_file)
    
    print("\n完成！")
    print(f"一般資訊替換: {result_stats['paragraphs'] + result_stats['tables']} 處")
    print(f"報廢資訊移除: {result_stats['obsolete_removed']} 處")
    print(f"檔案已儲存至: {output_file}")
